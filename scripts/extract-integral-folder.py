"""Batch extract all .docx files in integral/ folder to readable .txt.
Output: integral-extracted/<relative-path>.txt
"""
import os
import sys
import io
from pathlib import Path
from docx import Document

# Force UTF-8 stdout (Windows console default cp1252 can't print emoji/checkmark)
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

ROOT = Path(__file__).parent.parent
SRC_DIR = ROOT / "integral"
OUT_DIR = ROOT / "integral-extracted"
OUT_DIR.mkdir(exist_ok=True)

# Find all .docx files
docx_files = list(SRC_DIR.rglob("*.docx"))
# Skip lock files
docx_files = [f for f in docx_files if not f.name.startswith("~$")]

print(f"Found {len(docx_files)} .docx files")

success = 0
errors = []
for docx_path in docx_files:
    try:
        rel = docx_path.relative_to(SRC_DIR)
        # flatten path: replace / with _
        flat_name = str(rel).replace("\\", "_").replace("/", "_")
        flat_name = flat_name.replace(".docx", ".txt")
        out_path = OUT_DIR / flat_name

        doc = Document(str(docx_path))
        out_lines = []
        out_lines.append(f"# Source: {rel}")
        out_lines.append("")

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Track heading style
                style = para.style.name if para.style else ""
                if "Heading" in style:
                    level = "".join([c for c in style if c.isdigit()]) or "1"
                    out_lines.append("\n" + "#" * (int(level) + 1) + " " + text)
                else:
                    out_lines.append(text)

        for ti, table in enumerate(doc.tables):
            out_lines.append(f"\n--- TABLE {ti+1} ---")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                out_lines.append(" | ".join(cells))

        content = "\n".join(out_lines)
        out_path.write_text(content, encoding="utf-8")
        success += 1
        size = out_path.stat().st_size
        print(f"  ✓ {flat_name} ({size} bytes)")
    except Exception as e:
        errors.append((str(docx_path), str(e)))
        print(f"  ✗ {docx_path.name}: {e}", file=sys.stderr)

print(f"\nDone: {success}/{len(docx_files)} extracted, {len(errors)} errors")
if errors:
    for path, err in errors:
        print(f"  - {path}: {err}")
